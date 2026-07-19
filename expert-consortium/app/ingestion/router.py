"""Route each file to the right extractor and produce a Document ready for indexing."""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

from app.config import Domain
from app.ingestion import audio, ocr

TEXT_EXTS = {".txt", ".md"}


@dataclass
class Document:
    source_file: str
    text: str
    domain: Domain = "general"
    language: str = "unknown"
    ingested_at: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


def _extract_docx(path: Path) -> str:
    from docx import Document as DocxDocument  # imported lazily: heavy dependency

    doc = DocxDocument(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


EXTRACTORS: dict[str, Callable[[Path], str]] = {
    ".pdf": ocr.extract_pdf,
    **{ext: ocr.extract_image for ext in ocr.IMAGE_MIME},
    **{ext: audio.transcribe_audio for ext in audio.AUDIO_EXTS},
    **{ext: audio.extract_video for ext in audio.VIDEO_EXTS},
    ".docx": _extract_docx,
    **{ext: _extract_text for ext in TEXT_EXTS},
}

SUPPORTED_EXTS = frozenset(EXTRACTORS)


def guess_domain(path: Path) -> Domain:
    """Infer domain from the subfolder of uploads/ the file sits in, if any.

    uploads/law/contract.pdf -> "law".  Files directly in uploads/ -> "general".
    """
    parent = path.parent.name.lower()
    return parent if parent in ("law", "nko", "islamic", "cs") else "general"  # type: ignore[return-value]


def guess_language(text: str) -> str:
    """Cheap script-based language hint (enough for metadata filtering)."""
    sample = text[:4000]
    counts = {"ar": 0, "nko": 0, "latin": 0}
    for ch in sample:
        code = ord(ch)
        if 0x0600 <= code <= 0x06FF or 0x0750 <= code <= 0x077F:
            counts["ar"] += 1
        elif 0x07C0 <= code <= 0x07FF:
            counts["nko"] += 1
        elif ch.isalpha():
            counts["latin"] += 1
    best = max(counts, key=counts.get)  # type: ignore[arg-type]
    return best if counts[best] > 20 else "unknown"


def extract(path: Path, domain: Domain | None = None) -> Document:
    """Extract text from any supported file. Raises ValueError on unsupported types."""
    ext = path.suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(
            f"Unsupported file type '{ext}' ({path.name}). "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTS))}"
        )
    text = extractor(path)
    if not text.strip():
        raise ValueError(f"No text could be extracted from {path.name}")
    return Document(
        source_file=path.name,
        text=text,
        domain=domain or guess_domain(path),
        language=guess_language(text),
    )
